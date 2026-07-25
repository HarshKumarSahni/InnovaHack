# from openai import OpenAI
import google.generativeai as generativeai
import random
import time
import logging
import textwrap
import os
import pandas as pd
from docx import Document
from fpdf import FPDF
from services.PromptsDict import prompt_templates
from datetime import datetime
import json
import uuid
# import winsound
from pymongo import MongoClient, errors
from pymongo.errors import ConnectionFailure, OperationFailure
from typing import List, Dict, Any, Optional
import certifi
import cloudinary
import cloudinary.uploader

# ===============================================================
# === ROBUST PATH CONFIGURATION ===
# ===============================================================

# Get the absolute path of the directory containing this script (Generation.py)
# This will be .../src/MockTestAutomation/
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Get the project root directory by going up two levels from the script's directory
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, '..', '..'))

# Build absolute paths to ensure files are always found
MODEL = os.getenv("GENIE_MODEL", "gemini-2.5-flash-lite")
# Ensure only Gemini models are used; coerce/validate to avoid non-gemini model strings
if MODEL and "gemini" not in MODEL.lower():
    logging.warning("GENIE_MODEL appears to be non-Gemini ('%s'); defaulting to 'gemini-2.5-flash-lite'", MODEL)
    MODEL = "gemini-2.5-flash-lite"
SAVE_GENERATIONS_TO_DB = True
# Prefer environment variable for API key
API_KEY = os.getenv("GOOGLE_API_KEY") or os.getenv("API_KEY")
if API_KEY:
    # configure logging about which model/key is in use (do not log key itself)
    logging.basicConfig(level=logging.INFO)
    logging.info("Generative model set to: %s", MODEL)
else:
    # don't fail at import time; validation occurs before calls
    logging.basicConfig(level=logging.INFO)
    logging.warning("GOOGLE_API_KEY not set. Generation calls will raise until configured.")
# questions_per_chunk = 5 
# questions_per_chunk = 3 # change to 5 when on production level

# EXCEL_PATH = os.path.join(SCRIPT_DIR, "..", "data", "Syllabus.xlsx")
# EXCEL_PATH = os.path.join(SCRIPT_DIR, "..", "data", "UGCSyllabus.xlsx")

# Define the path to the backend/data directory
# BACKEND_DATA_DIR = os.path.join(PROJECT_ROOT, "backend", "data")
BACKEND_DATA_DIR = "/app/data"
# Create the output directories inside backend/data
OUTPUT_DIR = os.path.join(BACKEND_DATA_DIR, "generated_files")
RAW_RESPONSES_DIR = os.path.join(BACKEND_DATA_DIR, "raw_responses")

# Ensure the output directories exist at the project root
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(RAW_RESPONSES_DIR, exist_ok=True)

# ===============================================================
# === MONGODB SETUP ===
# ===============================================================
try:
    MONGO_CONNECTION_STRING = os.getenv("MONGO_URI")
    if not MONGO_CONNECTION_STRING:
        logging.warning("MONGO_URI not set. Logging disabled.")
        mongo_client = None
    else:
        # Added tlsAllowInvalidCertificates for local dev SSL issues
        # mongo_client = MongoClient(
        #     MONGO_CONNECTION_STRING,
        #     serverSelectionTimeoutMS=5000,
        #     tls=True,
        #     tlsCAFile=certifi.where(),
        #     # This is the "hammer" - it ignores certificate errors entirely
        #     tlsAllowInvalidCertificates=True,
        #     # Forces the client to connect even if the primary isn't immediately clear
        #     connectTimeoutMS=10000 
        # )
        mongo_client = MongoClient(
            MONGO_CONNECTION_STRING,
            serverSelectionTimeoutMS=5000,
            tls=True,
            tlsCAFile=certifi.where(), # Tells the container where the SSL certs are
            tlsAllowInvalidCertificates=True,
            connectTimeoutMS=10000 
        )
        # Verify connection
        mongo_client.admin.command('ping')
        logging.info("MongoDB connection successful.")
        db = mongo_client["acetrack_finetune_db"]
        finetune_collection = db["generation_logs"]
except Exception as e:
    logging.exception("MongoDB connection failed: %s", e)
    mongo_client = None
# ===============================================================
cloudinary.config(
    cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
    api_key=os.getenv("CLOUDINARY_API_KEY"),
    api_secret=os.getenv("CLOUDINARY_API_SECRET"),
    secure=True
)
# === TOPIC LOADING AND PROMPT GENERATION ===
# def load_all_topics(excel_path=EXCEL_PATH):
#     """Loads and shuffles topics from the specified Excel file."""
#     if not os.path.exists(excel_path):
#         raise FileNotFoundError(f"Syllabus file not found at the expected path: {excel_path}. Please ensure it exists.")
#     try:
#         df = pd.read_excel(excel_path)
#         if df.empty or len(df.columns) == 0:
#             raise ValueError("Syllabus.xlsx is empty or has no columns.")
#         topic_column = df.columns[0]
#         topics = df[topic_column].dropna().tolist()
#         random.shuffle(topics)
#         return topics
#     except Exception as e:
#         # Catch other potential pandas errors
#         raise RuntimeError(f"Failed to read or process the Excel file at {excel_path}: {e}")

def validate_topic_capacity(plan, total_topics, questions_per_chunk: int):
    """Validates if there are enough topics for the requested questions."""
    total_chunks_requested = sum(num // questions_per_chunk for num in plan.values())
    if total_chunks_requested > len(total_topics) // questions_per_chunk:
        error_msg = f"Not enough unique topics to generate the requested number of questions. \nTopics available: {len(total_topics)}, Questions requested: {sum(plan.values())}"
        raise ValueError(error_msg)

def build_prompt_from_template(topics_list, template_key, num_of_questions, EXAM):
    """Builds a Gemini prompt from a template with the given topics."""
    topics_str = "\n".join([f"{i+1}. {topic}" for i, topic in enumerate(topics_list)])
    randomized_answer_key = ', '.join(str(n) for n in random.choices(range(1, 5), k=num_of_questions))
    template = prompt_templates.get(template_key, "")
    return template.format(topics=topics_str, answer_key=randomized_answer_key, num=num_of_questions, exam=EXAM)


def build_live_quiz_prompt(topics_list, template_key, num_of_questions, EXAM):
    """Builds a Gemini prompt for generating a live quiz in JSON format."""
    topics_str = "\n".join([f"{i+1}. {topic}" for i, topic in enumerate(topics_list)])
    template = prompt_templates.get(template_key, "")
    return template.format(topics=topics_str, num=num_of_questions, exam=EXAM)


def parse_quiz_response(raw_text: str) -> List[Dict[str, Any]]:
    """Parses raw Gemini response into structured quiz questions."""
    try:
        data = json.loads(raw_text)
        questions = data.get("questions")
        if not isinstance(questions, list):
            raise ValueError("Parsed quiz output is missing questions array.")
        validated_questions = []
        for idx, item in enumerate(questions):
            if not isinstance(item, dict):
                raise ValueError(f"Question at index {idx} is not an object.")
            required_keys = {"id", "question", "options", "correct_answer", "explanation"}
            if not required_keys.issubset(set(item.keys())):
                raise ValueError(f"Question at index {idx} is missing required fields.")
            if not isinstance(item["options"], list):
                raise ValueError(f"The options field for question {idx} is not a list.")
            validated_questions.append(item)
        return validated_questions
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse quiz JSON response: {str(e)}")


def generate_all_prompts(plan, topics, exam, questions_per_chunk: int, live_mode: bool = False):
    """Generates a list of all prompts to be sent to the Gemini API."""
    prompts = []
    topic_index = 0
    shuffled_topics = random.sample(topics, len(topics))

    for qtype, count in plan.items():
        if topic_index + ((count // questions_per_chunk) * questions_per_chunk) > len(shuffled_topics):
            raise ValueError("Topic index out of bounds. This indicates a logic error in topic validation.")
        for _ in range(count // questions_per_chunk):
            chunk = shuffled_topics[topic_index: topic_index + questions_per_chunk]
            topic_index += questions_per_chunk
            if live_mode:
                prompt = build_live_quiz_prompt(chunk, qtype, questions_per_chunk, exam)
            else:
                prompt = build_prompt_from_template(chunk, qtype, questions_per_chunk, exam)
            prompts.append((qtype, prompt))
    return prompts


def extract_json_object(raw_text: str) -> str:
    """Attempts to isolate a JSON object from the raw Gemini text output."""
    if not raw_text or "{" not in raw_text:
        return raw_text
    start = raw_text.find("{")
    end = raw_text.rfind("}")
    if start >= 0 and end > start:
        return raw_text[start:end+1]
    return raw_text


def generate_live_quiz_task(topics: List[str], exam_name: str, question_count: int, testing_mode: bool):
    """Generates a live quiz as structured JSON data."""
    if not topics:
        raise ValueError("At least one topic is required for a live quiz.")
    live_template = prompt_templates.get("LIVE_QUIZ")
    if not live_template:
        raise ValueError("Live quiz prompt template is not configured.")

    formatted_topics = "\n".join([f"{i+1}. {topic}" for i, topic in enumerate(topics)])
    prompt = live_template.format(topics=formatted_topics, num=question_count, exam=exam_name)

    response_content, system_prompt = call_gemini(prompt, testing_mode, exam_name, 1)
    if not testing_mode:
        save_raw_response(response_content)
        log_generation_to_db(
            system_prompt=system_prompt,
            user_prompt=prompt,
            response_content=response_content,
            exam_name=exam_name,
            model_name=MODEL,
            testing=testing_mode,
        )

    cleaned_response = extract_json_object(response_content)
    questions = parse_quiz_response(cleaned_response)
    return {
        "success": True,
        "message": f"Generated {len(questions)} live quiz questions.",
        "quiz": questions
    }


def normalize_answer(answer: Optional[str]) -> str:
    return (answer or "").strip().upper()


def build_feedback_summary(results: List[Dict[str, Any]], score_percent: float, exam_name: str) -> Dict[str, Any]:
    strengths = []
    gaps = []
    for item in results:
        if item.get("is_correct"):
            strengths.append(
                f"Question {item['id']} was answered correctly with option {item['selected']}.")
        else:
            gaps.append(
                f"Question {item['id']} needs review: correct answer was {item['correct_answer']}, you selected {item['selected']}.")

    if not strengths:
        strengths = ["You showed good effort but should review the concepts behind the questions answered incorrectly."]
    if not gaps:
        gaps = ["No major gaps detected. Keep practicing at this difficulty level to reinforce your strengths."]

    resources = [
        "Review the explanations for any incorrect answers and revisit those topics in your study notes.",
        "Use timed practice quizzes to improve speed and accuracy under exam-like conditions.",
    ]
    if score_percent < 70:
        resources.append(
            "Focus on core concepts and fundamentals before attempting another high-difficulty quiz.")
    else:
        resources.append(
            "Continue with similar or slightly higher difficulty quizzes to push your performance further.")

    next_steps = (
        f"Your score is {score_percent}%. "
        + ("Maintain this pace and practice similarly difficult questions to build mastery."
           if score_percent >= 70 else
           "Review the areas above and practice targeted concept questions before retaking a similar quiz.")
    )

    return {
        "summary": (
            f"You completed the {exam_name} practice quiz with a score of {score_percent}%.")
            if exam_name else
            f"You completed the practice quiz with a score of {score_percent}%.",
        "score_percent": score_percent,
        "strengths": strengths,
        "gaps": gaps,
        "resources": resources,
        "next_steps": next_steps
    }


def grade_live_quiz_submission(quiz_questions: List[Dict[str, Any]], answers: Dict[str, str], exam_name: str):
    """Generates structured feedback for a completed quiz submission."""
    if not isinstance(quiz_questions, list):
        raise ValueError("Quiz questions must be a list.")
    if not isinstance(answers, dict):
        raise ValueError("Answers must be provided as a dictionary.")

    evaluated = []
    correct_count = 0
    for question in quiz_questions:
        qid = question.get("id") or question.get("question", "unknown").strip()[:32]
        correct_answer = normalize_answer(question.get("correct_answer"))
        selected = normalize_answer(answers.get(qid))
        is_correct = selected == correct_answer and correct_answer != ""
        evaluated.append({
            "id": qid,
            "question": question.get("question"),
            "selected": selected,
            "correct_answer": correct_answer,
            "is_correct": is_correct,
            "explanation": question.get("explanation")
        })
        if is_correct:
            correct_count += 1

    total_questions = len(quiz_questions)
    score_percent = round((correct_count / total_questions) * 100, 2) if total_questions > 0 else 0.0
    feedback_data = build_feedback_summary(evaluated, score_percent, exam_name)
    feedback_data.update({
        "correct_count": correct_count,
        "total_questions": total_questions,
    })
    return {
        "success": True,
        "message": "Quiz feedback generated successfully.",
        "feedback": feedback_data
    }

# === FILE OPERATIONS ===
def save_to_docx(content, filename):
    """Saves the given content to a .docx file in the output directory."""
    path = os.path.join(OUTPUT_DIR, filename)
    try:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(path)
    except Exception as e:
        raise IOError(f"❌ Cannot access {path}. Details: {e}")
    
def save_to_pdf(content, filename):
    """Saves the given content to a .pdf file in the output directory."""
    path = os.path.join(OUTPUT_DIR, filename)
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        encoded_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, txt=encoded_content)
        pdf.output(path)
    except Exception as e:
        raise IOError(f"❌ Cannot save PDF to {path}. Details: {e}")

def save_raw_response(text, folder=RAW_RESPONSES_DIR):
    """Saves the raw Gemini response for debugging purposes."""
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_id = uuid.uuid4().hex[:6]
    # filename = f"gemini_response_{timestamp}.docx"
    filename = f"gemini_response_{timestamp}_{file_id}.pdf"
    path = os.path.join(folder, filename)
    # doc = Document()
    # doc.add_paragraph(text)
    # doc.save(path)
    # print(f"✅ Raw response saved to: {path}")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        encoded_text = text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, txt=encoded_text)
        pdf.output(path)
        logging.info("Raw response saved to: %s", path)
    except Exception as e:
        logging.exception("Failed to save raw response PDF. Details: %s", e)

def log_generation_to_db(system_prompt: str, user_prompt: str, response_content: str, exam_name: str, model_name: str, testing: bool):
    """Logs a successful prompt/response pair to MongoDB if enabled."""
    if not testing and mongo_client and SAVE_GENERATIONS_TO_DB and response_content:
        try:
            finetune_collection.insert_one({
                "system": system_prompt,
                "prompt": user_prompt,
                "response": response_content,
                "exam": exam_name,
                "model": model_name,
                "created_at": datetime.now()
            })
            logging.info("Logged generation pair to MongoDB.")
        except OperationFailure as e:
            logging.warning("Failed to log to MongoDB (OperationFailure): %s", getattr(e, 'details', str(e)))
        except Exception as e:
            logging.exception("Failed to log to MongoDB (General Error): %s", e)
    elif not SAVE_GENERATIONS_TO_DB:
        logging.info("Skipping MongoDB log (master flag is OFF).")

# === Gemini handling ===
def _extract_response_text(response) -> str:
    """Robustly extract textual output from a Gemini/generativeai response object.

    The SDK returns different shapes across versions; try common attributes then
    fall back to stringifying the response.
    """
    try:
        # Older style: response.output[0].content[0].text
        if hasattr(response, "output") and response.output:
            first_output = response.output[0]
            content = getattr(first_output, "content", None)
            if content:
                if isinstance(content, (list, tuple)) and len(content) > 0:
                    item = content[0]
                    text = getattr(item, "text", None)
                    if text:
                        return text
                    if isinstance(item, dict):
                        return item.get("text") or json.dumps(item)
                elif isinstance(content, str):
                    return content

        # Newer SDK: response.candidates or response.generations
        if hasattr(response, "candidates") and response.candidates:
            cand = response.candidates[0]
            # candidate may be object or dict
            text = getattr(cand, "content", None) or getattr(cand, "text", None)
            if isinstance(text, str):
                return text
            if isinstance(text, (list, tuple)) and len(text) > 0 and isinstance(text[0], str):
                return text[0]
            if isinstance(cand, dict):
                return cand.get("content") or cand.get("text") or json.dumps(cand)

        # Fallbacks
        if hasattr(response, "text") and isinstance(response.text, str):
            return response.text
        # Last resort
        return str(response)
    except Exception:
        logging.exception("Failed to extract text from Gemini response")
        return str(response)


def _is_fatal_gemini_error(err_str: str) -> bool:
    """Return True if the error message indicates a non-transient error we should not retry."""
    s = (err_str or "").lower()
    # common fatal conditions
    fatal_keywords = [
        "invalid", "api key", "authentication", "not authorized", "permission denied",
        "model not found", "model doesn't exist", "model not available", "invalid model",
        "400", "400 bad request"
    ]
    return any(k in s for k in fatal_keywords)


def call_gemini(prompt, testing, exam_name, chunks, retries=3):
    """Calls the Google Gemini (Generative AI) API with retries and robust error handling.

    Returns (response_content, system_prompt) on success.
    """
    if testing:
        time.sleep(1)
        simulated_questions = []
        for i in range(chunks):
            simulated_questions.append({
                "id": f"q{i+1}",
                "question": f"Sample question {i+1} based on prompt type {prompt[:3]}",
                "options": ["A", "B", "C", "D"],
                "correct_answer": "A",
                "explanation": "This is a sample explanation."
            })
        return json.dumps({"questions": simulated_questions}), f"You are a {exam_name} paper setter."

    # Validate API key early
    if not API_KEY:
        raise RuntimeError("GOOGLE_API_KEY environment variable is not configured.")

    # configure SDK with the API key
    try:
        generativeai.configure(api_key=API_KEY)
        model_obj = generativeai.GenerativeModel(MODEL)
    except Exception as e:
        logging.exception("Failed to configure GenerativeModel: %s", e)
        raise

    system_prompt = f"You are a {exam_name} paper setter."
    full_prompt = f"{system_prompt}\n\n{prompt}"

    for attempt in range(retries):
        try:
            logging.info("Calling Gemini model '%s' attempt %d", MODEL, attempt + 1)
            # Prefer the modern `generate_content` API, but fall back to other methods on the model
            if hasattr(model_obj, 'generate_content'):
                response = model_obj.generate_content(
                    full_prompt,
                    generation_config={
                        "temperature": 0.7,
                        "max_output_tokens": 3000
                    }
                )
            elif hasattr(model_obj, 'generate'):
                # Older model objects may expose `generate` instead of `generate_content`
                response = model_obj.generate(
                    full_prompt,
                    temperature=0.7,
                    max_output_tokens=3000
                )
            elif hasattr(model_obj, 'predict'):
                response = model_obj.predict(full_prompt)
            else:
                raise RuntimeError("GenerativeModel instance has no recognized generation method")

            response_content = _extract_response_text(response)
            return response_content, system_prompt

        except Exception as e:
            err_str = str(e)
            logging.exception("Gemini attempt %d failed: %s", attempt + 1, err_str)
            # Decide whether to retry
            if _is_fatal_gemini_error(err_str):
                logging.error("Fatal Gemini error detected, will not retry: %s", err_str)
                raise
            # transient: wait and retry if attempts remain
            if attempt < retries - 1:
                sleep_time = 2 ** attempt
                logging.info("Retrying after %s seconds...", sleep_time)
                time.sleep(sleep_time)
                continue
            # exhausted
            logging.error("All Gemini retries exhausted")
            raise RuntimeError("All Gemini API retries failed.")

# Note: do not expose legacy `call_gpt` alias — use `call_gemini` explicitly.


# === CORE EXECUTION LOGIC ===
def handle_generation(prompts, TESTING, exam_name, questions_per_chunk: int):
    """Handles the question generation loop, calling Gemini for each prompt."""
    all_questions = []
    skipped_chunks = []
    # questions_per_chunk = 5
    max_retries_per_chunk = 3
    
    for qtype, prompt in prompts:
        logging.info("Generating questions for type: %s", qtype)
        generated_chunk = None
        last_failed_chunk = []
        response = None
        system_prompt_used = None
        
        for attempt in range(max_retries_per_chunk):
            try:
                logging.info("Attempt %d for %s...", attempt + 1, qtype)
                response, system_prompt_used = call_gemini(prompt, TESTING, exam_name, questions_per_chunk)
                logging.info("========== RAW GEMINI ==========")
                logging.info(response)
                logging.info("===============================")
                
                if response is None: # Handle potential failure from call_gemini retries
                    logging.error("call_gemini failed for %s after all retries.", qtype)
                    last_failed_chunk = [f"--- GEMINI CALL FAILED ---", f"Prompt Type: {qtype}"]
                    continue # Move to the next attempt or fail the chunk
                
                if not TESTING:
                    save_raw_response(response) 

                questions = [q.strip() for q in textwrap.dedent(response).split("--Question Starting--") if q.strip()]
                last_failed_chunk = questions
                
            #     if len(questions) != questions_per_chunk:
            #         print(f"⚠️ Gemini returned {len(questions)} questions instead of {questions_per_chunk}. Skipping this chunk.")
            #         skipped_chunks.append(questions)
            #         continue
                
            #     all_questions.extend(questions)
            # except Exception as e:
            #     print(f"An error occurred during generation for prompt type {qtype}: {e}")
            #     continue
                # --- VALIDATION LOGIC ---
                if len(questions) == questions_per_chunk:
                    logging.info("Success: got %d questions for %s.", len(questions), qtype)
                    generated_chunk = questions
                    if not TESTING and system_prompt_used: # Ensure system_prompt is available
                         log_generation_to_db(
                             system_prompt=system_prompt_used,
                             user_prompt=prompt, # Use the original user prompt
                             response_content=response, # Log the full raw response
                             exam_name=exam_name,
                             model_name=MODEL,
                             testing=TESTING
                         )
                    break # <<-- Exit the retry loop on success
                else:
                    logging.warning("Validation failed: Gemini returned %d questions instead of %d. Retrying...", len(questions), questions_per_chunk)
                    time.sleep(1) # Optional: wait a moment before retrying

            except Exception as e:
                logging.exception("An error occurred during Gemini call for %s: %s", qtype, e)
                if attempt < max_retries_per_chunk - 1:
                    time.sleep(2) # Wait longer if there's an actual API error
        
        # After the retry loop, check if we got a valid chunk
        if generated_chunk:
            all_questions.extend(generated_chunk)
        else:
            logging.error("Failed to generate a valid chunk for %s after %d attempts. Skipping.", qtype, max_retries_per_chunk)
            # Optionally, you could save the last failed response for debugging
            skipped_chunks.append(last_failed_chunk)
            
    # random.shuffle(all_questions)
    return all_questions, skipped_chunks


# === MAIN ENTRY POINT FOR BACKEND ===
def run_generation_task(plan: dict, testing_mode: bool, exam_name: str, output_format: str, questions_per_chunk: int, topics: List[str]):
    """Main function to be called by the FastAPI """
    try:
        logging.info("Starting generation for %s with plan: %s", exam_name, plan)
        
        run_id = uuid.uuid4().hex[:8]
        # questions_filename = f"Questions_{run_id}.docx"
        # skipped_filename = f"Skipped_{run_id}.docx"
        extension = ".pdf" if output_format == 'pdf' else ".docx"
        questions_filename = f"Questions_{run_id}{extension}"
        skipped_filename = f"Skipped_{run_id}{extension}"
        
        save_function = save_to_pdf if output_format == 'pdf' else save_to_docx

        # topics = load_all_topics()
        
        validate_topic_capacity(plan, topics, questions_per_chunk)
        
        prompts = generate_all_prompts(plan, topics, exam_name, questions_per_chunk)
        
        generated_questions, skipped_chunks = handle_generation(prompts, testing_mode, exam_name, questions_per_chunk)
        if not generated_questions and not skipped_chunks:
            raise RuntimeError("No questions were successfully generated. Check logs for API errors or response format issues.")
        
        # save_to_docx("\n\n".join(generated_questions), questions_filename)
        # save_to_pdf("\n\n".join(generated_questions), questions_filename)
        generated_files = {}
        message = ""
        if generated_questions:
            save_function("\n\n".join(generated_questions), questions_filename)
            # --- CLOUDINARY UPLOAD: Questions ---
            try:
                logging.info("Uploading %s to Cloudinary...", questions_filename)
                upload_questions = cloudinary.uploader.upload(
                    os.path.join(OUTPUT_DIR, questions_filename),
                    resource_type="raw",
                    public_id=f"ace-track/{questions_filename}"
                )
                # Store the Cloudinary URL instead of the local filename
                generated_files["questions"] = upload_questions.get("secure_url")
                message = f"Successfully generated {len(generated_questions)} questions."
            except Exception as u_err:
                logging.exception("Cloudinary upload failed for questions: %s", u_err)
                generated_files["questions"] = questions_filename # Fallback to filename
            
        if skipped_chunks:
            skipped_text = "\n\n".join([
                f"--- Skipped Chunk {i+1} ---:\n" + "\n\n".join(chunk)
                for i, chunk in enumerate(skipped_chunks)
            ])
            # save_to_docx(skipped_text, skipped_filename)
            save_function(skipped_text, skipped_filename)
            # --- CLOUDINARY UPLOAD: Skipped ---
            try:
                upload_skipped = cloudinary.uploader.upload(
                    os.path.join(OUTPUT_DIR, skipped_filename),
                    resource_type="raw",
                    public_id=f"ace-track/{skipped_filename}"
                )
                generated_files["skipped"] = upload_skipped.get("secure_url")
            except Exception as u_err:
                logging.exception("Cloudinary upload failed for skipped chunks: %s", u_err)
                generated_files["skipped"] = skipped_filename # Fallback
            
            
        if message and len(skipped_chunks)>0: # Add to existing message
                message += f" Failed to generate {len(skipped_chunks)} chunk(s), which have been saved separately."
        elif message:
            pass
        else: # Create new message
                message = f"Failed to generate questions, but {len(skipped_chunks)} skipped chunk(s) were saved."

        logging.info("Mock Test Generation Completed.")
        
        # if os.name == 'nt': # 'nt' is Windows
        #     try:
        #         import winsound
        #         winsound.PlaySound("CorrectHarp.wav", winsound.SND_FILENAME)
        #     except Exception as e:
        #         print(f"🟡 Windows notification sound failed: {e}")
        # else:
        #     # In Linux/Container environments, we just log completion
        #     print("🔔 Generation Task Finished")
        
        # generated_files = {"questions": questions_filename}
        # message = "Questions generated successfully."
        # if skipped_chunks:
        #     generated_files["skipped"] = skipped_filename

        return {
            "success": True,
            "message": message,
            "files": generated_files
            # "partial_success": bool(skipped_chunks and generated_questions) # True only if we have both
        }

    except Exception as e:
        error_message = f"Question generation failed: {str(e)}"
        logging.exception("%s", error_message)
        return {"success": False, "message": error_message, "files": {}}